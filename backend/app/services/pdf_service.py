import re
import logging
from dataclasses import dataclass, field

import pymupdf as fitz  # PyMuPDF (replaces deprecated `import fitz`)

logger = logging.getLogger(__name__)


@dataclass
class PageContent:
    """Extracted content from a single document page/section."""
    page_number: int
    text: str
    char_count: int
    word_count: int = 0
    headings: list[str] = field(default_factory=list)
    sections: list[str] = field(default_factory=list)
    paragraphs: list[str] = field(default_factory=list)


@dataclass
class PDFDocument:
    """Structured representation of an extracted document."""
    filename: str
    page_count: int
    full_text: str
    pages: list[PageContent] = field(default_factory=list)
    total_chars: int = 0
    total_words: int = 0
    headings: list[str] = field(default_factory=list)


class PDFServiceError(Exception):
    """Raised when document extraction encounters an error."""
    pass


class PDFService:
    """Service for extracting structured text and metadata from PDF, DOCX, TXT, and MD files."""

    def _extract_headings(self, text: str) -> list[str]:
        """Extract markdown style or capitalization headings from text."""
        headings = []
        for line in text.splitlines():
            line_str = line.strip()
            if line_str.startswith("#"):
                headings.append(line_str.lstrip("#").strip())
            elif line_str.isupper() and 3 < len(line_str) < 60:
                headings.append(line_str)
        return headings

    def _perform_ocr_fallback(self, doc: fitz.Document) -> list[PageContent]:
        """Perform OCR on pages with zero extractable text using pytesseract if available."""
        pages = []
        try:
            import pytesseract
            from PIL import Image
            import io

            logger.info("Attempting OCR fallback for image-only PDF pages...")
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                pix = page.get_pixmap(dpi=150)
                img = Image.open(io.BytesIO(pix.tobytes()))
                ocr_text = pytesseract.image_to_string(img).strip()
                
                words = ocr_text.split()
                paragraphs = [p.strip() for p in ocr_text.split("\n\n") if p.strip()]
                pages.append(PageContent(
                    page_number=page_num + 1,
                    text=ocr_text,
                    char_count=len(ocr_text),
                    word_count=len(words),
                    headings=self._extract_headings(ocr_text),
                    paragraphs=paragraphs
                ))
        except Exception as e:
            logger.warning(f"OCR fallback unavailable or failed: {e}")
        return pages

    def extract_from_bytes(self, file_bytes: bytes, filename: str) -> PDFDocument:
        """
        Extract text and metadata from PDF, DOCX, TXT, or MD raw bytes.
        """
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if ext == "docx":
            return self._extract_docx(file_bytes, filename)
        elif ext in ("txt", "md"):
            return self._extract_text(file_bytes, filename)
        else:
            return self._extract_pdf(file_bytes, filename)

    def _extract_pdf(self, file_bytes: bytes, filename: str) -> PDFDocument:
        """Extract content from PDF bytes using PyMuPDF and OCR fallback."""
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception as e:
            logger.error(f"Failed to open PDF '{filename}': {e}")
            raise PDFServiceError(
                f"Could not open '{filename}'. The file may be corrupt or password-protected."
            ) from e

        pages: list[PageContent] = []
        full_text_parts: list[str] = []
        all_headings: list[str] = []
        total_words = 0

        try:
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text = page.get_text("text").strip()
                words = text.split()
                paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
                headings = self._extract_headings(text)

                pages.append(PageContent(
                    page_number=page_num + 1,
                    text=text,
                    char_count=len(text),
                    word_count=len(words),
                    headings=headings,
                    paragraphs=paragraphs
                ))
                full_text_parts.append(text)
                all_headings.extend(headings)
                total_words += len(words)

            # Check if text density is zero (scanned PDF) and perform OCR
            if total_words == 0 and doc.page_count > 0:
                ocr_pages = self._perform_ocr_fallback(doc)
                if ocr_pages:
                    pages = ocr_pages
                    full_text_parts = [p.text for p in pages]
                    total_words = sum(p.word_count for p in pages)
                    all_headings = [h for p in pages for h in p.headings]

        except Exception as e:
            logger.error(f"Error extracting text from '{filename}': {e}")
            raise PDFServiceError(f"Failed to extract text from '{filename}'.") from e
        finally:
            doc.close()

        full_text = "\n\n".join(full_text_parts)

        return PDFDocument(
            filename=filename,
            page_count=len(pages),
            full_text=full_text,
            pages=pages,
            total_chars=len(full_text),
            total_words=total_words,
            headings=list(dict.fromkeys(all_headings))
        )

    def _extract_docx(self, file_bytes: bytes, filename: str) -> PDFDocument:
        """Extract content from DOCX file bytes."""
        try:
            import io
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)
            words = full_text.split()
            headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]

            page = PageContent(
                page_number=1,
                text=full_text,
                char_count=len(full_text),
                word_count=len(words),
                headings=headings,
                paragraphs=paragraphs
            )

            return PDFDocument(
                filename=filename,
                page_count=1,
                full_text=full_text,
                pages=[page],
                total_chars=len(full_text),
                total_words=len(words),
                headings=headings
            )
        except Exception as e:
            logger.error(f"Failed to parse DOCX file '{filename}': {e}")
            # Fallback to plain text decoding
            return self._extract_text(file_bytes, filename)

    def _extract_text(self, file_bytes: bytes, filename: str) -> PDFDocument:
        """Extract content from raw TXT or Markdown bytes."""
        try:
            text = file_bytes.decode("utf-8", errors="ignore").strip()
            words = text.split()
            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
            headings = self._extract_headings(text)

            page = PageContent(
                page_number=1,
                text=text,
                char_count=len(text),
                word_count=len(words),
                headings=headings,
                paragraphs=paragraphs
            )

            return PDFDocument(
                filename=filename,
                page_count=1,
                full_text=text,
                pages=[page],
                total_chars=len(text),
                total_words=len(words),
                headings=headings
            )
        except Exception as e:
            logger.error(f"Failed to parse text file '{filename}': {e}")
            raise PDFServiceError(f"Could not read text content from '{filename}'.") from e


def get_pdf_service() -> PDFService:
    """Dependency provider for PDFService."""
    return PDFService()
